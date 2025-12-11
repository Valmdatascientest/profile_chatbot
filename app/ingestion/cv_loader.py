"""
app/ingestion/cv_loader.py

Module pour charger un CV (PDF / DOCX) et le transformer
en texte structuré utilisable par le chatbot de profil.
"""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import List, Optional

import logging
import re
import textwrap

try:
    import pdfplumber
except ImportError:  # pragma: no cover - dépendance optionnelle
    pdfplumber = None

try:
    import docx  # python-docx
except ImportError:  # pragma: no cover - dépendance optionnelle
    docx = None

logger = logging.getLogger(__name__)


# =========================
#   Modèles de données
# =========================

@dataclass
class CVSection:
    title: str
    content: str
    section_type: Optional[str] = None  # ex: "experience", "education", "skills"


@dataclass
class CVDocument:
    path: Path
    raw_text: str
    sections: List[CVSection]


# =========================
#   Lecture du fichier
# =========================

def load_cv_text(cv_path: str | Path) -> str:
    """
    Charge le texte brut du CV à partir d'un fichier PDF ou DOCX.

    :param cv_path: chemin vers le CV (ex: data/raw/cv.pdf)
    :return: texte complet du CV
    """
    cv_path = Path(cv_path)
    if not cv_path.exists():
        raise FileNotFoundError(f"CV non trouvé: {cv_path}")

    suffix = cv_path.suffix.lower()

    if suffix == ".pdf":
        if pdfplumber is None:
            raise ImportError(
                "pdfplumber n'est pas installé. Installez-le avec `pip install pdfplumber`."
            )
        return _load_pdf_text(cv_path)

    if suffix in {".docx", ".doc"}:
        if docx is None:
            raise ImportError(
                "python-docx n'est pas installé. Installez-le avec `pip install python-docx`."
            )
        return _load_docx_text(cv_path)

    raise ValueError(f"Format de CV non supporté: {suffix}. Utilisez PDF ou DOCX.")


def _load_pdf_text(path: Path) -> str:
    texts: List[str] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            texts.append(page_text)
    full_text = "\n".join(texts)
    logger.info("CV PDF chargé: %s (%.1f kB)", path, path.stat().st_size / 1024)
    return full_text


def _load_docx_text(path: Path) -> str:
    document = docx.Document(str(path))
    texts: List[str] = []
    for para in document.paragraphs:
        text = para.text.strip()
        if text:
            texts.append(text)
    full_text = "\n".join(texts)
    logger.info("CV DOCX chargé: %s (%.1f kB)", path, path.stat().st_size / 1024)
    return full_text


# =========================
#   Découpage en sections
# =========================

# On définit quelques headings courants en FR/EN pour essayer de structurer le CV
SECTION_PATTERNS = {
    "experience": [
        r"^expériences? professionnelles?$",
        r"^experience professionnelle$",
        r"^work experience$",
        r"^professional experience$",
    ],
    "education": [
        r"^formation$",
        r"^éducation$",
        r"^education$",
        r"^studies$",
    ],
    "skills": [
        r"^compétences$",
        r"^skills$",
        r"^technical skills$",
        r"^compétences techniques$",
    ],
    "projects": [
        r"^projets?$",
        r"^projects?$",
        r"^personal projects$",
    ],
    "summary": [
        r"^profil$",
        r"^profil professionnel$",
        r"^résumé$",
        r"^summary$",
        r"^about me$",
        r"^à propos$",
    ],
}


def _guess_section_type(title: str) -> Optional[str]:
    title_norm = title.strip().lower()
    for section_type, patterns in SECTION_PATTERNS.items():
        for p in patterns:
            if re.match(p, title_norm, flags=re.IGNORECASE):
                return section_type
    return None


def split_cv_into_sections(raw_text: str) -> List[CVSection]:
    """
    Essaye de découper le CV en sections à partir des titres
    (Expérience, Formation, Compétences, etc.).

    Si aucun titre n'est détecté, renvoie une seule section globale.
    """
    lines = [l.strip() for l in raw_text.splitlines()]
    lines = [l for l in lines if l]  # supprime les lignes vides

    if not lines:
        return []

    # Heuristique simple : une ligne en majuscules ou très courte
    # peut être un titre de section.
    sections: List[CVSection] = []
    current_title = "CV"
    current_lines: List[str] = []

    def flush_section():
        nonlocal current_title, current_lines, sections
        if current_lines:
            content = "\n".join(current_lines).strip()
            if content:
                section_type = _guess_section_type(current_title)
                sections.append(
                    CVSection(
                        title=current_title.strip(),
                        content=content,
                        section_type=section_type,
                    )
                )
        current_lines = []

    for line in lines:
        is_heading = _is_potential_heading(line)
        if is_heading:
            # On termine la section précédente
            flush_section()
            current_title = line
        else:
            current_lines.append(line)

    # Dernière section
    flush_section()

    # Si on a qu'une seule section "CV" et pas de structure, ce n'est pas grave
    if not sections:
        sections.append(
            CVSection(
                title="CV",
                content=raw_text.strip(),
                section_type=None,
            )
        )

    return sections


def _is_potential_heading(line: str) -> bool:
    """
    Heuristique pour détecter un titre de section :

    - ligne courte (<= 5 mots)
    - ET (majoritairement en majuscules OU correspond à un pattern connu)
    """
    if len(line) > 80:
        return False

    # Si on matche déjà les patterns connus, c'est un heading
    if _guess_section_type(line) is not None:
        return True

    words = line.split()
    if len(words) <= 5:
        # ratio majuscules
        uppercase_chars = sum(c.isupper() for c in line if c.isalpha())
        alpha_chars = sum(c.isalpha() for c in line)
        if alpha_chars > 0 and uppercase_chars / alpha_chars > 0.6:
            return True

    return False


# =========================
#   Conversion en objet & chunks
# =========================

def load_cv_document(cv_path: str | Path) -> CVDocument:
    """
    Charge un CV et le découpe en sections.

    :param cv_path: chemin du CV (pdf/docx)
    :return: CVDocument
    """
    cv_path = Path(cv_path)
    raw_text = load_cv_text(cv_path)
    sections = split_cv_into_sections(raw_text)
    return CVDocument(path=cv_path, raw_text=raw_text, sections=sections)


def cv_to_text_chunks(
    cv: CVDocument,
    max_chars: int = 800,
    min_chars: int = 200,
) -> List[str]:
    """
    Transforme un CVDocument en liste de chunks de texte de taille raisonnable
    pour les embeddings (RAG).

    :param cv: CVDocument déjà chargé
    :param max_chars: taille max d'un chunk
    :param min_chars: taille minimale avant de fusionner
    """
    chunks: List[str] = []

    for section in cv.sections:
        section_header = f"[Section: {section.title}]"
        base_text = f"{section_header}\n{section.content}"

        if len(base_text) <= max_chars:
            chunks.append(base_text)
        else:
            # on découpe la section en sous-chunks
            parts = _split_text_into_chunks(base_text, max_chars=max_chars)
            chunks.extend(parts)

    # Fusion simple de très petits chunks si nécessaire
    merged: List[str] = []
    buffer = ""

    for chunk in chunks:
        if not buffer:
            buffer = chunk
        elif len(buffer) + len(chunk) < min_chars:
            buffer = buffer + "\n\n" + chunk
        else:
            merged.append(buffer)
            buffer = chunk

    if buffer:
        merged.append(buffer)

    return merged


def _split_text_into_chunks(text: str, max_chars: int = 800) -> List[str]:
    """
    Découpe un texte en morceaux d'au plus max_chars.
    Essaye de couper sur les sauts de ligne ou les phrases.
    """
    # On commence par couper sur les double sauts de ligne si possible
    paragraphs = re.split(r"\n\s*\n", text)
    chunks: List[str] = []

    for para in paragraphs:
        para = para.strip()
        if not para:
            continue
        if len(para) <= max_chars:
            chunks.append(para)
        else:
            # On fait un découpage plus grossier par largeur
            # (ça reste suffisant pour un CV)
            wrapped = textwrap.wrap(para, width=max_chars, break_long_words=False)
            chunks.extend(wrapped)

    return chunks


if __name__ == "__main__":
    # Petit test manuel : python -m app.ingestion.cv_loader data/raw/cv.pdf
    import sys

    logging.basicConfig(level=logging.INFO)

    if len(sys.argv) < 2:
        print("Usage: python -m app.ingestion.cv_loader <path_to_cv>")
        raise SystemExit(1)

    path = Path(sys.argv[1])
    cv_doc = load_cv_document(path)
    print(f"Sections détectées: {len(cv_doc.sections)}")
    for s in cv_doc.sections:
        print("====")
        print(f"Titre: {s.title!r} | type: {s.section_type}")
        print(s.content[:300], "...")
        print()

    chunks = cv_to_text_chunks(cv_doc)
    print(f"\nNombre de chunks générés: {len(chunks)}")
    for c in chunks[:3]:
        print("----")
        print(c[:400])
        print()